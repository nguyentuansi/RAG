"""Async document ingestion pipeline with multi-format support."""

from __future__ import annotations

import asyncio
import io
import mimetypes
import time
from collections.abc import AsyncIterator, Callable
from pathlib import Path
from typing import Any
from uuid import uuid4

import httpx

from src.rag.core.exceptions import UnsupportedDocumentFormatError
from src.rag.core.logging import get_logger
from src.rag.domain.documents import Document, DocumentFormat, DocumentMetadata

logger = get_logger(__name__)


class DocumentIngestionPipeline:
    """
    Async ingestion pipeline supporting PDF, DOCX, HTML, Markdown, and plain text.

    All I/O operations are non-blocking; CPU-bound parsing (PDF, DOCX) is
    dispatched to the default ThreadPoolExecutor.
    """

    def __init__(
        self,
        max_file_size_mb: int = 50,
        http_timeout: float = 30.0,
    ) -> None:
        self._max_bytes = max_file_size_mb * 1024 * 1024
        self._http_timeout = http_timeout
        self._loop = asyncio.get_event_loop

    # ------------------------------------------------------------------ #
    #  Public interface                                                    #
    # ------------------------------------------------------------------ #

    async def ingest_file(self, path: Path) -> Document:
        """Ingest a file from the local filesystem."""
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        stat = path.stat()
        if stat.st_size > self._max_bytes:
            raise ValueError(
                f"File {path.name} ({stat.st_size // 1024} KB) exceeds limit of "
                f"{self._max_bytes // 1024 // 1024} MB"
            )

        fmt = DocumentFormat.from_extension(path.suffix)
        raw_bytes = await asyncio.get_running_loop().run_in_executor(
            None, path.read_bytes
        )
        content = await self._parse(raw_bytes, fmt)
        return Document(
            content=content,
            format=fmt,
            metadata=DocumentMetadata(source=str(path), title=path.name),
        )

    async def ingest_url(self, url: str, *, title: str | None = None) -> Document:
        """Fetch and ingest a URL. Supports HTML and plain text responses."""
        async with httpx.AsyncClient(timeout=self._http_timeout) as client:
            response = await client.get(url, follow_redirects=True)
            response.raise_for_status()

        content_type = response.headers.get("content-type", "text/html").split(";")[0]
        fmt = _mime_to_format(content_type)
        content = await self._parse(response.content, fmt)

        return Document(
            content=content,
            format=fmt,
            metadata=DocumentMetadata(
                source=url,
                title=title or url.split("/")[-1],
            ),
        )

    async def ingest_batch(
        self,
        sources: list[Path | str],
        *,
        on_progress: Callable[[int, int], None] | None = None,
    ) -> list[Document]:
        """Ingest multiple sources concurrently with optional progress callback."""
        docs: list[Document] = []
        tasks = [
            self.ingest_file(Path(s)) if isinstance(s, (Path, str)) and not str(s).startswith("http")
            else self.ingest_url(str(s))
            for s in sources
        ]

        for i, coro in enumerate(asyncio.as_completed(tasks)):
            try:
                doc = await coro
                docs.append(doc)
            except Exception as exc:
                logger.error("batch_ingest_item_failed", index=i, error=str(exc))

            if on_progress:
                on_progress(i + 1, len(tasks))

        return docs

    # ------------------------------------------------------------------ #
    #  Format parsers                                                      #
    # ------------------------------------------------------------------ #

    async def _parse(self, raw: bytes, fmt: DocumentFormat) -> str:
        loop = asyncio.get_running_loop()
        parser = {
            DocumentFormat.PDF: self._parse_pdf,
            DocumentFormat.DOCX: self._parse_docx,
            DocumentFormat.HTML: self._parse_html,
            DocumentFormat.MARKDOWN: self._parse_text,
            DocumentFormat.TXT: self._parse_text,
            DocumentFormat.JSON: self._parse_text,
            DocumentFormat.CSV: self._parse_text,
        }.get(fmt)

        if parser is None:
            raise UnsupportedDocumentFormatError(f"No parser for format: {fmt}")

        if fmt in (DocumentFormat.PDF, DocumentFormat.DOCX):
            return await loop.run_in_executor(None, parser, raw)
        return parser(raw)

    def _parse_pdf(self, raw: bytes) -> str:
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(raw))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(p for p in pages if p.strip())
        except ImportError:
            raise UnsupportedDocumentFormatError(
                "pypdf not installed. Run: pip install pypdf"
            )

    def _parse_docx(self, raw: bytes) -> str:
        try:
            import docx

            doc = docx.Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            raise UnsupportedDocumentFormatError(
                "python-docx not installed. Run: pip install python-docx"
            )

    def _parse_html(self, raw: bytes) -> str:
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(raw, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            return raw.decode("utf-8", errors="replace")

    def _parse_text(self, raw: bytes) -> str:
        return raw.decode("utf-8", errors="replace")


def _mime_to_format(mime: str) -> DocumentFormat:
    mapping = {
        "text/html": DocumentFormat.HTML,
        "application/pdf": DocumentFormat.PDF,
        "text/plain": DocumentFormat.TXT,
        "text/markdown": DocumentFormat.MARKDOWN,
        "application/json": DocumentFormat.JSON,
        "text/csv": DocumentFormat.CSV,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentFormat.DOCX,
    }
    return mapping.get(mime, DocumentFormat.TXT)
